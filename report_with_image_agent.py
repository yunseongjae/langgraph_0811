from typing import Annotated, TypedDict, List, Dict
from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages

# State definition
class State(TypedDict):
    messages: Annotated[list, add_messages]
    outline: Dict[str, str]
    current_section: int
    section_content: str
    section_image: str
    image_prompt: str
    total_sections: int
    full_report: List[Dict[str, str]]

# Graph builder initialization
graph_builder = StateGraph(State)

from langchain_community.tools.tavily_search import TavilySearchResults

search = TavilySearchResults(max_results=3)

from pydantic import BaseModel, Field, create_model

# Function to create dynamic Outline model
def create_outline_model(section_count: int):
    fields = {f"section{i}": 
              (str, Field(description=f"Title for section {i}")) for i in range(1, section_count + 1)}
    return create_model("DynamicOutline", **fields)

# Outline generation node
def outline_generator(state: State):
    DynamicOutline = create_outline_model(state["total_sections"])
    outline_parser = JsonOutputParser(pydantic_object=DynamicOutline)

    outline_prompt = PromptTemplate(
        template="Create an outline for a detailed report with exactly {section_count} main sections.\n{format_instructions}\nThe topic is: {topic}\n",
        input_variables=["section_count", "topic"],
        partial_variables={"format_instructions": outline_parser.get_format_instructions()},
    )
    
    chain = outline_prompt | llm | outline_parser
    
    try:
        outline = chain.invoke({"section_count": state["total_sections"], "topic": state["messages"][-1].content})
        return {"outline": outline, "current_section": 1, "full_report": []}
    except Exception as e:
        print(f"Error in create_outline: {e}")
        return {"error": str(e)}


from openai import OpenAI
from langchain_core.tools import tool

client = OpenAI()

# Schema definition for DALL-E image generation
class GenImageSchema(BaseModel):
    prompt: str = Field(description="The prompt for image generation")

# DALL-E image generation function definition
@tool(args_schema=GenImageSchema)
def generate_image(prompt: str) -> str:
    """Generate an image using DALL-E based on the given prompt."""
    if not prompt:
        raise ValueError("Image generation prompt cannot be empty")
    try:
        response = client.images.generate(
            model="dall-e-3",
            prompt=prompt,
            size="1024x1024",
            quality="standard",
            n=1
        )
        return response.data[0].url
    except Exception as e:
        print(f"Error in image generation: {e}")
        return "Image generation failed"

def image_generator(state: State):

    image_prompt = PromptTemplate(
        template="""Based on the following section content, 
        create a prompt for generating an infographic that represents this section.
        Section content: {section_content}
        \n\n
        Image generation prompt(under 500 characters):""",
        input_variables=["section_content"],
    )
    
    image_prompt = llm.invoke(image_prompt.format(section_content=state["section_content"]))
    image_url = generate_image(image_prompt.content)

    current_section = {
        "title": state['outline'][f"section{state['current_section']}"],
        "content": state['section_content'],
        "image_url": image_url,
        "image_prompt": "DO NOT CONTAIN ANY METRIC OR TEXT ON THE IMAGE. \n"+image_prompt.content if isinstance(image_prompt, AIMessage) else image_prompt
    }

    updated_full_report = state.get("full_report", []) + [current_section]

    print(f"Completed section {state['current_section']} of {state['total_sections']}")

    return {
        "image_prompt": image_prompt.content if isinstance(image_prompt, AIMessage) else image_prompt,
        "section_image": image_url,
        "current_section": state["current_section"] + 1,
        "full_report": updated_full_report
    }

from langchain_openai import ChatOpenAI

llm = ChatOpenAI(model="gpt-4o-mini")

from langchain_core.messages import HumanMessage, AIMessage
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate

def contents_writer(state: State):
    if "error" in state:
        return {"messages": [AIMessage(content=f"An error occurred: {state['error']}")]}
    
    if state["current_section"] > state["total_sections"]:
        return {"messages": [AIMessage(content="Report completed.")]}
        
    current_section_key = f"section{state['current_section']}"
    current_topic = state["outline"][current_section_key]
    search_results = search.invoke(current_topic)    
    
    previous_sections_content = []
    for i in range(1, state['current_section']):
        section_key = f"section{i}"
        if section_key in state["section_content"]:
            previous_sections_content.append(f"Section {i}: {state['outline'][section_key]}\n{state['section_content'][section_key]}")
    
    previous_sections = "\n\n".join(previous_sections_content)
    
    section_prompt = PromptTemplate(
        template="""Write a detailed section for the topic: {topic}. 
        
        Use the following search results for information: {search_results}
        
        Previous sections:
        {previous_sections}
        Write only the content for this section, 
        do not include any image prompts or suggestions.
        Detailed statistics or information is needed, 
        so you should include collected information from search result.""",
        input_variables=["topic", "search_results", "previous_sections"],
    )
    section_content = llm.invoke(section_prompt.format(
        topic=current_topic,
        search_results=search_results,
        previous_sections=previous_sections
    ))

    return {
        "section_content": section_content.content,
        "current_section": state["current_section"]
    }

from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO

def report_generator(state: State):
    doc = Document()
    doc.add_heading(f"Report: {state['messages'][0].content}", 0)

    for section in state['full_report']:
        doc.add_heading(section['title'], level=1)
        doc.add_paragraph(section['content'])
        
        # 이미지 추가
        if section['image_url'] != "Image generation failed":
            try:
                response = requests.get(section['image_url'])
                image = BytesIO(response.content)
                doc.add_picture(image, width=Inches(6))
                doc.add_paragraph(f"Image prompt: {section['image_prompt']}")
            except Exception as e:
                doc.add_paragraph(f"Failed to add image: {str(e)}")

        doc.add_page_break()

    # 보고서 저장
    filename = f"report_{state['messages'][0].content}.docx".replace(" ", "_")
    doc.save(filename)

    return {
        "messages": [AIMessage(content=f"Report finalized and saved as {filename}.")],
        "report_file": filename
    }
# Add nodes
graph_builder.add_node("outline_generator", outline_generator)
graph_builder.add_node("contents_writer", contents_writer)
graph_builder.add_node("image_generator", image_generator)
graph_builder.add_node("report_generator", report_generator)

# Add edges
graph_builder.add_edge(START, "outline_generator")
graph_builder.add_edge("outline_generator", "contents_writer")
graph_builder.add_edge("contents_writer", "image_generator")
graph_builder.add_edge("report_generator", END)
# Add conditional edges
def should_continue_writing(state: State):
    if state["current_section"] <= state["total_sections"]:
        return "write_section"
    else:
        return "finalize_report"

graph_builder.add_conditional_edges(
    "image_generator",
    should_continue_writing,
    {
        "write_section": "contents_writer",
        "finalize_report": "report_generator"
    }
)

# Compile graph
graph = graph_builder.compile()
